#!/usr/bin/env python3
"""
Convert HTML manual to DOCX format
"""

import os
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

def html_to_docx():
    """Convert HTML manual to DOCX format"""
    
    # Paths
    html_path = "docs/AirAd_Vendor_Portal_User_Manual.html"
    docx_path = "docs/AirAd_Vendor_Portal_User_Manual.docx"
    
    print("🚀 Converting HTML manual to DOCX...")
    print("=" * 50)
    
    # Check if HTML file exists
    if not os.path.exists(html_path):
        print(f"❌ HTML file not found: {html_path}")
        return False
    
    print(f"✓ HTML file found: {html_path}")
    
    # Read HTML file
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create Word document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Add title
    title = soup.find('title')
    if title:
        title_para = doc.add_heading(title.get_text(), 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Process content
    content = soup.find('body')
    if not content:
        print("❌ No body content found in HTML")
        return False
    
    # Remove script and style tags
    for script in content(["script", "style"]):
        script.decompose()
    
    # Process each element
    for element in content.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'div', 'section']):
        tag_name = element.name
        
        # Skip empty elements
        if not element.get_text(strip=True):
            continue
        
        if tag_name.startswith('h'):
            # Headings
            level = int(tag_name[1])
            heading = doc.add_heading(element.get_text(strip=True), level)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
        elif tag_name == 'p':
            # Paragraphs
            para = doc.add_paragraph()
            para.text = element.get_text(strip=True)
            
        elif tag_name in ['ul', 'ol']:
            # Lists
            list_items = element.find_all('li', recursive=False)
            for li in list_items:
                text = li.get_text(strip=True)
                if text:
                    p = doc.add_paragraph(text, style='List Bullet')
                    
        elif tag_name == 'div':
            # Handle divs with specific classes
            if 'page-break' in element.get('class', []):
                # Add page break
                doc.add_page_break()
            else:
                # Treat as paragraph
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
    
    # Save document
    doc.save(docx_path)
    
    # Get file size
    file_size = os.path.getsize(docx_path)
    
    print(f"✓ DOCX created successfully: {docx_path}")
    print(f"  File size: {file_size:,} bytes")
    
    return True

def main():
    """Main function"""
    try:
        success = html_to_docx()
        if success:
            print("\n🎉 DOCX conversion completed successfully!")
            print("\n📁 Available manual formats:")
            print("  📄 AirAd_Vendor_Portal_User_Manual.html")
            print("  📄 AirAd_Vendor_Portal_User_Manual.pdf")
            print("  📄 AirAd_Vendor_Portal_User_Manual.docx")
            print("  📄 AirAd_Vendor_Portal_User_Manual.txt")
            print("\n📂 Manual location: docs/")
            print("✅ Ready for distribution to vendors!")
        else:
            print("\n❌ DOCX conversion failed!")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())
