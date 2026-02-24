"""Sections 5-9: Geo, Vendors, CSV Import, Google Places Seeding, Field Ops, QA"""
from docx.shared import Pt
from ug_helpers import h1, h2, body, bullet, note_box, table, step, pb


def build(doc):

    # ── SECTION 5 ─────────────────────────────────────────────────────────
    h1(doc, "5.  Managing Locations (Geo Hierarchy)")
    body(doc, "Before adding any vendors, you must set up the geographic structure of the platform. Think of it as nested boxes: Country → City → Area → Landmark.")
    table(doc, ["Level", "Example", "Description"], [
        ["Country",  "Pakistan",        "Top-level geographic unit"],
        ["City",     "Karachi",         "A city within a country"],
        ["Area",     "Defence Phase 5", "A neighbourhood or district within a city"],
        ["Landmark", "Dolmen Mall",     "A specific well-known location within an area (optional)"],
    ])
    body(doc, "Every vendor must be assigned to a City and an Area. Landmark is optional but helps with precise AR map placement.")
    h2(doc, "5.1  How to Add a New City")
    step(doc, 1, "Go to the Geo section", "Click 'Geo' in the left navigation menu.")
    step(doc, 2, "Click 'Add City'", "Find and click the Add City button.")
    step(doc, 3, "Fill in the details", "Enter the required information:", subs=[
        "Name: Full city name (e.g. 'Karachi')",
        "Slug: Short, lowercase, hyphen-separated identifier (e.g. 'karachi'). CANNOT be changed later.",
        "Country: Select from the dropdown",
        "Display Order: Controls order in dropdowns (lower number = appears first)",
        "Active: Tick to make the city available in the system",
    ])
    step(doc, 4, "Save", "Click Save. The city appears in the list.")
    note_box(doc, "IMPORTANT", "The 'slug' is permanent. Once set, it cannot be changed. Use only lowercase letters, numbers, and hyphens (e.g. 'karachi', 'lahore', 'new-york').")
    h2(doc, "5.2  How to Add a New Area")
    step(doc, 1, "Go to Geo", "Click 'Geo' in the navigation.")
    step(doc, 2, "Select the city", "Choose which city this area belongs to.")
    step(doc, 3, "Click 'Add Area'", "Click the Add Area button.")
    step(doc, 4, "Fill in details", "Enter area name, slug, and optionally a parent area for nested zones.")
    step(doc, 5, "Save", "Click Save.")
    h2(doc, "5.3  How to Add a Landmark")
    step(doc, 1, "Go to Geo", "Navigate to the Geo section.")
    step(doc, 2, "Select the area", "Choose the area this landmark is in.")
    step(doc, 3, "Click 'Add Landmark'", "Click the Add Landmark button.")
    step(doc, 4, "Fill in details", "Enter landmark name, slug, and GPS coordinates (latitude and longitude).")
    step(doc, 5, "Save", "Click Save.")
    note_box(doc, "WARNING", "GPS coordinates must be real, valid numbers. Latitude: -90 to 90. Longitude: -180 to 180. The system rejects 0,0 as invalid.")
    pb(doc)

    # ── SECTION 6 ─────────────────────────────────────────────────────────
    h1(doc, "6.  Vendors — The Heart of the Portal")
    body(doc, "A Vendor is any business or point of interest on the AirAd platform — a restaurant, shop, clinic, school, market stall, etc. Adding and managing vendors is the primary task for most users.")
    h2(doc, "6.1  The Vendor List Page")
    body(doc, "Click 'Vendors' in the navigation to see all vendors. You can search and filter by:")
    table(doc, ["Filter", "What It Does"], [
        ["Search by name",       "Type any part of a business name to find it"],
        ["Filter by City",       "Show only vendors in a specific city"],
        ["Filter by Area",       "Show only vendors in a specific area"],
        ["Filter by QC Status",  "Show vendors with a specific review status (PENDING, APPROVED, etc.)"],
        ["Filter by Data Source","Show vendors added via CSV, manual entry, or field agents"],
        ["Filter by Tag",        "Show vendors with a specific tag assigned"],
    ])
    h2(doc, "6.2  How to Add a New Vendor (Manual Entry)")
    body(doc, "Roles that can add vendors: Super Admin, City Manager, Data Entry.")
    step(doc, 1, "Go to Vendors", "Click 'Vendors' in the navigation menu.")
    step(doc, 2, "Click 'Add Vendor'", "Find and click the Add Vendor or '+' button.")
    step(doc, 3, "Fill in required fields", "These fields must be completed:", subs=[
        "Business Name: Official trading name (e.g. 'Al-Fatah Superstore')",
        "Slug: Unique URL-friendly identifier — auto-generated but editable",
        "City: Select from dropdown",
        "Area: Select the area within the chosen city",
        "GPS Latitude: Exact latitude of the business (e.g. 24.8607)",
        "GPS Longitude: Exact longitude of the business (e.g. 67.0011)",
    ])
    step(doc, 4, "Fill in optional fields", "These improve data quality:", subs=[
        "Description: Short description of the business",
        "Address Text: Human-readable street address",
        "Phone Number: Business phone (stored securely, encrypted — never visible in plain text)",
        "Landmark: Nearest landmark (select from dropdown)",
        "Business Hours: Opening/closing times for each day of the week",
        "Data Source: How data was collected (Manual Entry, Field Agent, CSV Import, Google Places)",
    ])
    step(doc, 5, "Save the vendor", "Click 'Save' or 'Create Vendor'. Status is automatically set to PENDING.")
    note_box(doc, "IMPORTANT", "After saving, the vendor status is PENDING. It will NOT appear on the AirAd app until a QA Reviewer approves it.")
    h2(doc, "6.3  How to Edit a Vendor")
    step(doc, 1, "Find the vendor", "Use the search or filters on the Vendor List page.")
    step(doc, 2, "Click on the vendor", "Click the vendor name or row to open the detail page.")
    step(doc, 3, "Click 'Edit'", "Click the Edit button.")
    step(doc, 4, "Make your changes", "Update any fields that need correction.")
    step(doc, 5, "Save", "Click Save. All changes are recorded in the Audit Log automatically.")
    h2(doc, "6.4  How to Delete a Vendor")
    body(doc, "Roles that can delete vendors: Super Admin, City Manager.")
    note_box(doc, "IMPORTANT", "Deleting a vendor does NOT permanently remove it from the database. It is 'soft deleted' — hidden from the app but preserved for audit purposes.")
    step(doc, 1, "Find the vendor", "Search for the vendor on the Vendor List page.")
    step(doc, 2, "Open the vendor detail", "Click on the vendor.")
    step(doc, 3, "Click 'Delete'", "Click the Delete button.")
    step(doc, 4, "Confirm", "Confirm the deletion. The vendor is hidden from the app but remains in the database.")
    h2(doc, "6.5  Vendor Data Source Types")
    table(doc, ["Data Source", "Meaning"], [
        ["Manual Entry",  "Added directly by a Data Entry user through the portal"],
        ["CSV Import",    "Added via a bulk CSV file upload"],
        ["Field Agent",   "Added or verified by a Field Agent during an on-site visit"],
        ["Google Places", "Imported from Google Places API (automated)"],
    ])
    h2(doc, "6.6  Vendor Photos")
    body(doc, "Each vendor can have photos attached. Photos are uploaded during Field Visits (see Section 8). You can view a vendor's photos from the Vendor Detail page under the 'Photos' tab.")
    note_box(doc, "NOTE", "Photos are stored securely. The system generates temporary secure links when you view them — they cannot be shared as permanent public links.")
    h2(doc, "6.7  Vendor Tags")
    body(doc, "Tags are labels that categorise vendors. You can assign up to 15 tags per vendor. See Section 10 for full details on tags.")
    note_box(doc, "IMPORTANT", "A vendor must have at least one active CATEGORY tag assigned before it can be approved by a QA Reviewer. Without a category tag, the approval will be blocked.")
    h2(doc, "6.8  Understanding Business Hours")
    body(doc, "Business hours can be set for each day of the week. For each day you specify:")
    bullet(doc, "Whether the business is open on that day (Yes/No)", bold_prefix="Open")
    bullet(doc, "When the business opens (e.g. 09:00)", bold_prefix="Opening Time")
    bullet(doc, "When the business closes (e.g. 22:00)", bold_prefix="Closing Time")
    note_box(doc, "TIP", "If a business is open 24 hours, set opening time to 00:00 and closing time to 23:59. If closed on a specific day, toggle that day to 'Closed'.")
    pb(doc)

    # ── SECTION 7 ─────────────────────────────────────────────────────────
    h1(doc, "7.  Importing Vendors (CSV Upload & Google Places Seeding)")
    body(doc, "The Import Center provides two ways to add vendors in bulk. CSV Upload lets you upload a spreadsheet file. Google Places Seeding lets you pull business listings directly from Google. Both methods are available from the same page.")
    h2(doc, "7.1  Who Can Import?")
    table(doc, ["Import Method", "Who Can Use It"], [
        ["CSV Upload",           "Super Admin, City Manager, Data Entry"],
        ["Google Places Seeding","Super Admin, City Manager, Operations Manager"],
    ])
    body(doc, "If your role has access to both methods, you will see two tabs at the top of the Import Center page: 'CSV upload' and 'Google Places'. Click the tab you need. If your role only has access to one method, you will only see that method — no tabs are shown.")
    h2(doc, "7.2  What is a CSV File?")
    body(doc, "A CSV (Comma-Separated Values) file is a simple spreadsheet format that can be created in Microsoft Excel, Google Sheets, or LibreOffice Calc. It stores data in rows and columns, with each row representing one business.")
    h2(doc, "7.3  Required CSV Columns")
    body(doc, "Your CSV file MUST have these exact column headers in the first row:")
    table(doc, ["Column Name", "Required?", "Description", "Example Value"], [
        ["business_name", "YES", "The official name of the business",                    "Al-Fatah Superstore"],
        ["latitude",      "YES", "GPS latitude of the business location",                "24.8607"],
        ["longitude",     "YES", "GPS longitude of the business location",               "67.0011"],
        ["city_slug",     "YES", "The slug of the city (must already exist in system)",  "karachi"],
        ["area_slug",     "YES", "The slug of the area (must already exist in system)",  "defence-phase-5"],
        ["phone",         "No",  "Business phone number",                                "+92-21-1234567"],
        ["description",   "No",  "Short description of the business",                   "A large supermarket chain"],
        ["address_text",  "No",  "Human-readable street address",                       "Plot 5, Main Boulevard, DHA"],
    ])
    note_box(doc, "WARNING", "Column names are case-sensitive. Use exactly: business_name, latitude, longitude, city_slug, area_slug. Any typo will cause the entire import to fail.")
    h2(doc, "7.4  How to Prepare Your CSV File")
    step(doc, 1, "Open Excel or Google Sheets", "Create a new spreadsheet.")
    step(doc, 2, "Add the column headers", "In Row 1, type the column names exactly as shown in the table above.")
    step(doc, 3, "Add your data", "From Row 2 onwards, enter one business per row.")
    step(doc, 4, "Check GPS coordinates", "Ensure all latitude/longitude values are real numbers. Do not leave them blank or enter 0.", subs=[
        "Latitude must be between -90 and 90 (e.g. 24.8607 for Karachi)",
        "Longitude must be between -180 and 180 (e.g. 67.0011 for Karachi)",
        "Do NOT enter 0 for either — the system will reject it",
    ])
    step(doc, 5, "Check city_slug and area_slug", "These must exactly match slugs already set up in the Geo section. If a slug does not exist, that row will fail.", subs=[
        "Go to the Geo section to find the correct slugs",
        "Slugs are lowercase with hyphens (e.g. 'karachi', 'defence-phase-5')",
    ])
    step(doc, 6, "Save as CSV", "In Excel: File → Save As → choose 'CSV (Comma delimited)'. In Google Sheets: File → Download → CSV.")
    h2(doc, "7.5  How to Upload a CSV File")
    step(doc, 1, "Go to Imports", "Click 'Imports' in the left navigation menu.")
    step(doc, 2, "Click 'New Import'", "Find and click the New Import or Upload button.")
    step(doc, 3, "Select your CSV file", "Click 'Browse' or 'Choose File' and select your prepared CSV file from your computer.")
    step(doc, 4, "Click 'Upload'", "Click the Upload or Submit button.")
    step(doc, 5, "Wait for processing", "The system processes your file in the background. You will see the import appear in the list with status PROCESSING.")
    step(doc, 6, "Check the results", "Once processing is complete, the status changes to COMPLETED or FAILED. Click on the import to see details.")
    note_box(doc, "NOTE", "Large files may take a few minutes to process. Do not close the browser or upload the same file again while it is processing.")
    h2(doc, "7.6  Understanding Import Statuses")
    table(doc, ["Status", "Meaning", "What To Do"], [
        ["PENDING",                 "Upload received, waiting to start processing",          "Wait — processing will begin automatically"],
        ["PROCESSING",              "File is being read and vendors are being created",       "Wait — do not re-upload"],
        ["COMPLETED",               "All rows processed successfully",                        "Review the imported vendors in the Vendor list"],
        ["COMPLETED_WITH_ERRORS",   "Some rows succeeded, some failed",                       "Click the import to see which rows failed and why"],
        ["FAILED",                  "The entire import failed (e.g. wrong file format)",      "Fix the file and upload again"],
    ])
    h2(doc, "7.7  Understanding Import Error Reports")
    body(doc, "When an import has errors, click on it to see a detailed error report. Each error shows the row number and the reason it failed. Common errors include:")
    table(doc, ["Error Message", "What It Means", "How to Fix"], [
        ["Missing required column: business_name", "The column header is missing or misspelled",      "Check your CSV headers — must be exactly 'business_name'"],
        ["Row 5: city_slug 'xyz' not found",       "The city slug in that row does not exist",        "Check the Geo section for the correct slug and update your CSV"],
        ["Row 12: Invalid latitude '0'",           "The latitude value is 0 or invalid",              "Enter a real GPS coordinate for that business"],
        ["Row 8: Duplicate vendor",                "A vendor with the same name and location exists", "Check if the business is already in the system"],
        ["Row 20: area_slug not in city",          "The area does not belong to the specified city",  "Ensure the area_slug matches an area within the given city_slug"],
    ])
    note_box(doc, "TIP", "The error log shows a maximum of 1,000 errors. If your file has more than 1,000 errors, fix the first batch of errors and re-upload.")

    # ── SECTION 7.8+ — Google Places Seeding ────────────────────────────
    h2(doc, "7.8  What is Google Places Seeding?")
    body(doc, "Google Places Seeding is a way to automatically import business listings from Google into the AirAd system. Instead of typing each business by hand or preparing a spreadsheet, you select a geographic area and the system finds businesses in that area using Google's database.")
    body(doc, "This is especially useful for quickly populating a new city or neighbourhood with vendor data before field agents visit in person.")
    note_box(doc, "NOTE", "Businesses imported from Google Places are added with a status of 'Pending'. They must still be reviewed and approved by a QA Reviewer before they appear on the AirAd app.")

    h2(doc, "7.9  Who Can Use Google Places Seeding?")
    body(doc, "Roles that can seed from Google Places: Super Admin, City Manager, Operations Manager.")
    body(doc, "Data Entry users do NOT have access to Google Places Seeding. They can only use CSV Upload.")

    h2(doc, "7.10  How to Seed Vendors from Google Places")
    step(doc, 1, "Go to the Import Center", "Click 'Imports' in the left navigation menu.")
    step(doc, 2, "Switch to the Google Places tab", "If you see two tabs at the top, click 'Google Places'. If you only see the Google Places form (no tabs), you are already on the right page.")
    step(doc, 3, "Select a country", "Open the Country dropdown and choose the country where you want to find businesses. If no countries appear, ask your Super Admin to set up countries in the Geo section first.")
    step(doc, 4, "Select a city", "After choosing a country, the City dropdown becomes available. Select the city. Only cities belonging to the selected country are shown.")
    step(doc, 5, "Select an area", "After choosing a city, the Area dropdown becomes available. Select the specific neighbourhood or district. Only areas belonging to the selected city are shown.")
    step(doc, 6, "Set the search radius", "Enter a number between 100 and 5000. This is the distance in metres from the centre of the selected area. A larger number searches a wider area. The default is 1500 metres (1.5 kilometres).")
    step(doc, 7, "Enter a search query", "Type keywords describing the type of businesses you want to find. The default is 'restaurants food'. You can change this to anything — for example 'pharmacy', 'supermarket', 'school', or 'salon barbershop'.")
    step(doc, 8, "Select category tags (optional)", "Below the search fields, you will see a row of category tag chips. Click any tags that match the type of businesses you are importing. Selected tags turn orange. You can select multiple tags or leave them all unselected.", subs=[
        "Tags are optional — they help narrow the search and are automatically applied to imported vendors",
        "Click a selected tag again to deselect it",
    ])
    step(doc, 9, "Click 'Seed vendors'", "Click the orange 'Seed vendors' button in the bottom-right corner of the form. The system starts searching for businesses in the background.")
    step(doc, 10, "Watch the progress", "A new row appears in the 'Google Places batches' table below the form. The status changes from 'Pending' to 'In progress' and finally to 'Done'. The progress bar shows how many businesses have been processed.")
    note_box(doc, "NOTE", "You can close the browser or navigate to other pages while the import runs. It processes in the background. Come back to the Import Center to check progress at any time.")

    h2(doc, "7.11  Google Places Seed Form — Field by Field")
    table(doc, ["Field", "Required?", "What to Enter", "What Happens if Left Empty"], [
        ["Country",        "Yes", "Select the country from the dropdown",                     "You cannot proceed — City dropdown stays disabled"],
        ["City",           "Yes", "Select the city (appears after choosing Country)",          "You cannot proceed — Area dropdown stays disabled"],
        ["Area",           "Yes", "Select the area or neighbourhood",                         "The 'Seed vendors' button stays disabled"],
        ["Search radius",  "Yes", "A number between 100 and 5000 (metres)",                   "Shows a warning: 'Must be between 100 and 5000'"],
        ["Search query",   "Yes", "Keywords describing the businesses (e.g. 'restaurants')",   "The 'Seed vendors' button stays disabled"],
        ["Category tags",  "No",  "Click zero or more tag chips to select them",              "All business types are included — no filtering applied"],
    ])

    h2(doc, "7.12  Understanding the Google Places Batches Table")
    body(doc, "Below the seed form, a table shows all your Google Places import jobs. Each row represents one seeding request. The table columns are:")
    table(doc, ["Column", "What It Shows"], [
        ["Area",     "The area you selected when you started the seed"],
        ["Query",    "The search keywords you used"],
        ["Status",   "The current state of the import (see status table below)"],
        ["Progress", "A progress bar and counter showing how many businesses have been processed out of the total found"],
        ["Errors",   "The number of businesses that could not be imported (e.g. missing data from Google)"],
        ["Started",  "The date and time you started the seed"],
    ])
    body(doc, "The table refreshes automatically every 5 seconds while any import is still running. You can also click 'Refresh' to update manually.")

    h2(doc, "7.13  Google Places Seed Statuses")
    table(doc, ["Status", "What It Means", "What To Do"], [
        ["Pending",     "Your request has been received and is waiting to start",              "Wait — processing starts automatically within a few seconds"],
        ["In progress", "The system is actively searching Google and importing businesses",    "Wait — watch the progress bar. Do not submit the same area again"],
        ["Done",        "All businesses found have been processed",                            "Review the imported vendors in the Vendor list"],
        ["Failed",      "Something went wrong during processing",                              "The system will automatically retry up to 3 times. If it still fails, try again with a smaller radius or different keywords"],
    ])

    h2(doc, "7.14  Google Places — Duplicate Protection")
    body(doc, "The system has multiple layers of protection to prevent duplicate data:")
    bullet(doc, "If you try to seed the same area with the same keywords while a previous seed is still running, the system will show an orange warning: 'An identical import is already in progress'. Wait for the first one to finish.", bold_prefix="Same area, same query")
    bullet(doc, "If a business already exists in the AirAd system (from a previous import or manual entry), the system updates the existing record instead of creating a duplicate.", bold_prefix="Business already exists")
    bullet(doc, "If a seed job fails partway through, you can start it again. The system remembers which businesses it already processed and skips them, continuing from where it left off.", bold_prefix="Resumed after failure")
    note_box(doc, "TIP", "You do not need to worry about duplicates. The system handles them automatically. It is safe to seed the same area multiple times — existing businesses will be refreshed, not duplicated.")

    h2(doc, "7.15  Google Places — Common Error Messages")
    table(doc, ["Message", "What It Means", "How to Fix"], [
        ["An identical import is already in progress",              "You submitted the same area and keywords while a job is still running", "Wait for the running job to finish, then try again if needed"],
        ["Area has no centroid coordinates set",                    "The selected area does not have a centre point defined in the system",  "Ask your Super Admin or City Manager to set GPS coordinates for that area in the Geo section"],
        ["Please check your selections and try again",             "One of your selections is invalid (e.g. area does not belong to city)", "Make sure you select Country, then City, then Area in order"],
        ["Something went wrong. Please try again later",           "A temporary system error occurred",                                     "Wait a moment and try again. If the problem continues, contact your Super Admin"],
    ])

    h2(doc, "7.16  Google Places — Frequently Asked Questions")
    for q, a in [
        ("How many businesses will be imported?",
         "It depends on the area and search keywords. Google returns up to 60 results per search. Larger or busier areas with popular search terms will return more results."),
        ("Can I import all business types at once?",
         "Yes — leave the category tags empty and use a broad search query like 'business' or 'shop'. However, for better results, it is recommended to do focused searches like 'restaurants', 'pharmacy', 'supermarket' separately."),
        ("Will imported businesses appear on the AirAd app immediately?",
         "No. All imported businesses start with a status of Pending. A QA Reviewer must approve each one before it appears on the app."),
        ("What if a business has closed since Google last updated?",
         "Google data may not always be current. Field Agents should verify imported businesses with on-site visits. This is why all imports start as Pending."),
        ("Can I cancel a running import?",
         "Currently, you cannot cancel an import once it has started. It will run to completion. If it is importing unwanted data, you can review and reject those vendors afterwards."),
        ("Why does the area dropdown show no options?",
         "The selected city may not have any areas set up yet. Ask your Super Admin or City Manager to add areas for that city in the Geo section."),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f"Q: {q}")
        r.bold = True; r.font.size = Pt(10.5)
        doc.add_paragraph(f"A: {a}").runs[0].font.size = Pt(10.5)
        doc.add_paragraph()
    pb(doc)

    # ── SECTION 8 ─────────────────────────────────────────────────────────
    h1(doc, "8.  Field Operations — On-Site Visits & Photos")
    body(doc, "Field Operations cover the work done by Field Agents who physically visit business locations to verify data, collect photos, and confirm GPS accuracy.")
    h2(doc, "8.1  What is a Field Visit?")
    body(doc, "A Field Visit is a recorded on-site visit by a Field Agent to a vendor's physical location. It captures:")
    bullet(doc, "The date and time of the visit")
    bullet(doc, "The confirmed GPS location (latitude and longitude measured on-site)")
    bullet(doc, "Notes about the visit")
    bullet(doc, "Photos taken at the location")
    bullet(doc, "The visit outcome (e.g. Verified, Not Found, Closed Down)")
    h2(doc, "8.2  Who Can Record Field Visits?")
    body(doc, "Roles that can create field visits: Super Admin, City Manager, Field Agent.")
    body(doc, "Field Agents can only see their own visits. City Managers and Super Admins can see all visits.")
    h2(doc, "8.3  How to Record a Field Visit")
    step(doc, 1, "Go to Field Ops", "Click 'Field Ops' in the left navigation menu.")
    step(doc, 2, "Click 'New Visit'", "Find and click the New Visit or '+' button.")
    step(doc, 3, "Select the vendor", "Search for and select the vendor you are visiting.")
    step(doc, 4, "Enter visit details", "Fill in the required information:", subs=[
        "Visit Date: The date you visited (today's date by default)",
        "Confirmed GPS Latitude: The GPS reading from your device at the location",
        "Confirmed GPS Longitude: The GPS reading from your device at the location",
        "Outcome: Select the result (Verified, Not Found, Closed Down, Needs Update)",
        "Notes: Any additional observations about the business",
    ])
    step(doc, 5, "Save the visit", "Click Save. The visit is recorded and linked to the vendor.")
    note_box(doc, "TIP", "Use your phone's GPS or a GPS app to get accurate coordinates while standing at the business location. Do not guess or copy coordinates from the internet.")
    h2(doc, "8.4  How to Upload Photos for a Visit")
    step(doc, 1, "Open the field visit", "Go to Field Ops and click on the visit you want to add photos to.")
    step(doc, 2, "Click 'Add Photos'", "Find and click the Add Photos or Upload button.")
    step(doc, 3, "Select photos", "Choose one or more photos from your device.")
    step(doc, 4, "Upload", "Click Upload. Photos are saved and linked to the visit and vendor.")
    note_box(doc, "NOTE", "Photos are stored securely in the cloud. They are linked to both the field visit and the vendor record. You can view them from the vendor's detail page.")
    h2(doc, "8.5  Visit Outcomes Explained")
    table(doc, ["Outcome", "Meaning", "What Happens Next"], [
        ["Verified",      "Business exists and details are confirmed correct",   "Vendor data is confirmed; QA Reviewer may approve"],
        ["Not Found",     "No business found at the given location",             "Vendor should be reviewed and possibly deleted"],
        ["Closed Down",   "Business has permanently closed",                     "Vendor should be soft-deleted or flagged"],
        ["Needs Update",  "Business exists but some details need correction",    "Data Entry user should update the vendor record"],
    ])
    h2(doc, "8.6  GPS Drift Detection")
    body(doc, "The system automatically compares the GPS coordinates recorded during a field visit with the GPS coordinates stored for the vendor. If there is a large difference (GPS drift), the vendor is automatically flagged as 'Needs Review' for QA inspection.")
    note_box(doc, "NOTE", "GPS drift does not mean the data is wrong — it may simply mean the original coordinates were slightly off. A QA Reviewer will check and correct if needed.")
    pb(doc)

    # ── SECTION 9 ─────────────────────────────────────────────────────────
    h1(doc, "9.  Quality Assurance (QA) Review")
    body(doc, "Quality Assurance (QA) is the process of checking vendor data for accuracy before it goes live on the AirAd app. Every vendor must pass QA review before it is visible to app users.")
    h2(doc, "9.1  Who Does QA?")
    body(doc, "Roles that can perform QA reviews: Super Admin, City Manager, QA Reviewer.")
    h2(doc, "9.2  The QA Dashboard")
    body(doc, "The QA Dashboard shows all vendors that are currently flagged as 'Needs Review'. To access it:")
    step(doc, 1, "Go to QA", "Click 'QA' in the left navigation menu.")
    step(doc, 2, "View the list", "You will see all vendors requiring review, sorted by priority.")
    step(doc, 3, "Click a vendor", "Click on any vendor to open its full detail page for review.")
    h2(doc, "9.3  Understanding QC Statuses")
    table(doc, ["Status", "Meaning", "Who Sets It"], [
        ["PENDING",      "Newly added vendor, not yet reviewed",                        "Set automatically when a vendor is created"],
        ["APPROVED",     "Vendor has passed QA review and is live on the app",          "Set by QA Reviewer or City Manager"],
        ["REJECTED",     "Vendor failed QA review and will not appear on the app",      "Set by QA Reviewer or City Manager"],
        ["NEEDS_REVIEW", "Vendor has been flagged for re-inspection",                   "Set automatically (GPS drift, duplicate) or manually"],
        ["FLAGGED",      "Vendor has been manually flagged for attention",              "Set by any authorised user"],
    ])
    h2(doc, "9.4  How to Approve a Vendor")
    step(doc, 1, "Open the vendor", "Find the vendor in the QA Dashboard or Vendor List and click on it.")
    step(doc, 2, "Review all details", "Check the business name, GPS location, address, phone number, tags, and photos.")
    step(doc, 3, "Verify the vendor has a CATEGORY tag", "A vendor cannot be approved without at least one active Category tag.")
    step(doc, 4, "Click 'Approve'", "Click the Approve button.")
    step(doc, 5, "Confirm", "The vendor status changes to APPROVED and it becomes visible on the AirAd app.")
    note_box(doc, "IMPORTANT", "Once approved, a vendor is immediately visible on the AirAd app. Make sure all details are correct before approving.")
    h2(doc, "9.5  How to Reject a Vendor")
    step(doc, 1, "Open the vendor", "Find the vendor and click on it.")
    step(doc, 2, "Review the issues", "Identify what is wrong with the data (e.g. incorrect GPS, missing details, duplicate entry).")
    step(doc, 3, "Click 'Reject'", "Click the Reject button.")
    step(doc, 4, "Add a reason (optional)", "You may add a note explaining why the vendor was rejected.")
    step(doc, 5, "Confirm", "The vendor status changes to REJECTED and it will not appear on the app.")
    h2(doc, "9.6  How to Flag a Vendor for Review")
    body(doc, "Any authorised user can manually flag a vendor for QA review:")
    step(doc, 1, "Open the vendor", "Find and open the vendor detail page.")
    step(doc, 2, "Click 'Flag for Review'", "Click the Flag or Mark for Review button.")
    step(doc, 3, "Add a note", "Describe the issue you have noticed.")
    step(doc, 4, "Save", "The vendor status changes to NEEDS_REVIEW and appears in the QA Dashboard.")
    h2(doc, "9.7  Automatic QA Flags")
    body(doc, "The system automatically flags vendors in two situations:")
    bullet(doc, "GPS Drift: When a field visit confirms GPS coordinates that differ significantly from the stored coordinates", bold_prefix="GPS Drift Detection")
    bullet(doc, "When the system detects two or more vendors with very similar names and locations", bold_prefix="Duplicate Detection")
    note_box(doc, "NOTE", "Automatic flags are not errors — they are alerts for human review. A QA Reviewer will investigate and either approve, correct, or reject the vendor.")
    h2(doc, "9.8  Duplicate Detection")
    body(doc, "The system scans for potential duplicate vendors by comparing business names and GPS locations. If two vendors are found to be very similar:")
    bullet(doc, "Both vendors are flagged as NEEDS_REVIEW")
    bullet(doc, "They appear in the QA Dashboard side by side")
    bullet(doc, "A QA Reviewer decides which record to keep and which to reject")
    pb(doc)
