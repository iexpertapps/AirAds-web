"""Sections 1-4: Intro, Roles, Login, Dashboard"""
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ug_helpers import h1, h2, body, bullet, note_box, table, step, pb


def build(doc):
    # COVER
    doc.add_paragraph(); doc.add_paragraph()
    for txt, sz, col, bold in [
        ("AirAd Platform", 28, (21,67,96), True),
        ("Data Collection Portal", 22, (40,116,166), True),
        ("", 10, (0,0,0), False),
        ("Comprehensive User Guide", 18, (80,80,80), True),
        ("", 10, (0,0,0), False),
        ("A complete step-by-step guide for all portal users", 12, (100,100,100), False),
        ("No technical knowledge required", 12, (100,100,100), True),
        ("", 10, (0,0,0), False),
        ("Version 1.0  |  Phase A  |  Internal Use Only", 10, (130,130,130), False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.bold = bold; r.font.size = Pt(sz); r.font.color.rgb = RGBColor(*col)
    pb(doc)

    # TABLE OF CONTENTS
    h1(doc, "Table of Contents")
    for n, t in [
        ("1","Introduction & What is the Data Collection Portal?"),
        ("2","Who Uses This Portal? — User Roles Explained"),
        ("3","Getting Started — Logging In & Your Account"),
        ("4","The Dashboard — Your Home Screen"),
        ("5","Managing Locations (Geo Hierarchy)"),
        ("6","Vendors — The Heart of the Portal"),
        ("7","Importing Vendors (CSV Upload & Google Places Seeding)"),
        ("8","Field Operations — On-Site Visits & Photos"),
        ("9","Quality Assurance (QA) Review"),
        ("10","Tags — Organising & Categorising Vendors"),
        ("11","Governance — Fraud, Blacklists & Enforcement"),
        ("12","Audit Log — Full Activity History"),
        ("13","User Management (Super Admin Only)"),
        ("14","Understanding Statuses at a Glance"),
        ("15","Common Tasks — Quick Reference"),
        ("16","Troubleshooting & FAQs"),
        ("17","Glossary of Terms"),
        ("18","Data Privacy & Security"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.add_run(f"  {n}.  {t}").font.size = Pt(11)
    pb(doc)

    # ── SECTION 1 ─────────────────────────────────────────────────────────
    h1(doc, "1.  Introduction & What is the Data Collection Portal?")
    body(doc, "Welcome to the AirAd Data Collection Portal User Guide.")
    body(doc,
        "The AirAd Data Collection Portal is a secure, web-based administration system used by the "
        "AirAd team to collect, organise, verify, and manage business (vendor) data across multiple "
        "cities. This data powers the AirAd mobile application, which helps users discover local "
        "businesses using augmented reality (AR) technology.")
    body(doc,
        "Think of this portal as the 'back office' of AirAd. Everything that appears on the AirAd "
        "app — business names, locations, photos, categories — is first entered, checked, and "
        "approved here in this portal.")
    h2(doc, "1.1  What Can You Do in This Portal?")
    for t, d in [
        ("Add new businesses (vendors)", "Enter business details like name, address, GPS location, phone number, and opening hours."),
        ("Import businesses in bulk", "Upload a spreadsheet (CSV file) to add hundreds of businesses at once."),
        ("Verify business data", "Review and approve or reject business listings to ensure accuracy."),
        ("Manage locations", "Set up cities, areas, and landmarks that businesses belong to."),
        ("Organise with tags", "Label businesses with categories, promotions, and other tags."),
        ("Field operations", "Record on-site visits by field agents and upload photos."),
        ("Monitor quality", "Flag businesses that need review due to GPS errors or duplicate entries."),
        ("Governance & compliance", "Manage fraud scores, blacklists, and enforcement actions."),
        ("Audit trail", "View a complete history of every action taken in the system."),
    ]:
        bullet(doc, d, bold_prefix=t)
    note_box(doc, "NOTE", "This guide is written for non-technical users. No programming or IT knowledge is needed to use this portal.")
    pb(doc)

    # ── SECTION 2 ─────────────────────────────────────────────────────────
    h1(doc, "2.  Who Uses This Portal? — User Roles Explained")
    body(doc, "Every person who uses the portal is assigned a Role. Your role determines what you can see and what actions you can perform. There are 11 roles in total.")
    h2(doc, "2.1  Data Collection Roles")
    table(doc, ["Role", "Who They Are", "What They Can Do"], [
        ["Super Admin", "Highest-level administrator", "Full access — create users, manage all data, delete records, unlock accounts, view all reports"],
        ["City Manager", "Manages data for a specific city", "Add/edit/delete vendors, manage imports, view field visits, approve QA reviews"],
        ["Data Entry", "Enters business data", "Add and edit vendors, upload CSV imports, assign tags"],
        ["Operations Manager", "Oversees platform operations and data seeding", "Seed vendors from Google Places, view governance data, manage suspensions, view audit logs"],
        ["QA Reviewer", "Checks data quality", "Review and approve/reject vendor listings, view QA dashboard, flag issues"],
        ["Field Agent", "Visits businesses on-site", "Record field visits, upload photos, view own visits only"],
        ["Analyst", "Analyses platform data", "View vendors, view audit logs — read-only access"],
        ["Support", "Handles support queries", "View vendor listings — read-only"],
    ])
    h2(doc, "2.2  Governance & Specialist Roles")
    table(doc, ["Role", "Who They Are", "What They Can Do"], [
        ["Content Moderator", "Reviews content policy violations", "Access governance page, manage enforcement actions"],
        ["Data Quality Analyst", "Ensures data accuracy", "Manage tags, review vendor data quality"],
        ["Analytics Observer", "Monitors platform metrics", "View analytics and reports — read-only"],
    ])
    note_box(doc, "IMPORTANT", "You cannot change your own role. Only a Super Admin can assign or change roles. Contact your Super Admin if you believe your role is incorrect.")
    h2(doc, "2.3  Portal Access by Role — Quick Summary")
    table(doc, ["Portal Section", "Super Admin", "City Manager", "Data Entry", "Ops Manager", "QA Reviewer", "Field Agent", "Analyst"], [
        ["Dashboard",                "Yes", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["Vendors (View)",           "Yes", "Yes", "Yes", "No",  "Yes", "Yes", "Yes"],
        ["Vendors (Add/Edit)",       "Yes", "Yes", "Yes", "No",  "No",  "No",  "No"],
        ["Vendors (Delete)",         "Yes", "Yes", "No",  "No",  "No",  "No",  "No"],
        ["CSV Import",               "Yes", "Yes", "Yes", "No",  "No",  "No",  "No"],
        ["Google Places Seeding",    "Yes", "Yes", "No",  "Yes", "No",  "No",  "No"],
        ["Field Ops",                "Yes", "Yes", "No",  "No",  "No",  "Yes", "No"],
        ["QA Review",                "Yes", "Yes", "No",  "No",  "Yes", "No",  "No"],
        ["Geo Management",           "Yes", "Yes", "Yes", "No",  "No",  "No",  "No"],
        ["Tags",                     "Yes", "Yes", "Yes", "No",  "No",  "No",  "No"],
        ["Governance",               "Yes", "No",  "No",  "Yes", "No",  "No",  "No"],
        ["Audit Log",                "Yes", "No",  "No",  "Yes", "No",  "No",  "Yes"],
        ["User Management",          "Yes", "No",  "No",  "No",  "No",  "No",  "No"],
    ])
    pb(doc)

    # ── SECTION 3 ─────────────────────────────────────────────────────────
    h1(doc, "3.  Getting Started — Logging In & Your Account")
    h2(doc, "3.1  How to Log In")
    body(doc, "Follow these steps to access the portal:")
    step(doc, 1, "Open your web browser", "Use Google Chrome, Firefox, Edge, or Safari. The portal works best on a desktop or laptop.")
    step(doc, 2, "Go to the portal address", "Type the portal URL into your browser's address bar. Your Super Admin will provide the correct web address.")
    step(doc, 3, "Enter your email address", "Type the email address registered for your account (usually your work email).")
    step(doc, 4, "Enter your password", "Type your password. First-time users should use the temporary password provided by the Super Admin.")
    step(doc, 5, "Click 'Login'", "If your details are correct, you will be taken to the Dashboard.")
    note_box(doc, "WARNING", "If you see 'Account locked', too many incorrect password attempts were made. Wait 15 minutes and try again, or ask your Super Admin to unlock your account.")
    h2(doc, "3.2  Account Security & Password Tips")
    for b in [
        "Use at least 8 characters",
        "Mix uppercase and lowercase letters",
        "Include at least one number and one special character (e.g. !, @, #)",
        "Never share your password with anyone",
        "Do not use your name or email as your password",
        "Change your temporary password immediately on first login",
    ]:
        bullet(doc, b)
    h2(doc, "3.3  Account Lockout Policy")
    body(doc, "Your account locks automatically after 5 wrong password attempts in a row.")
    bullet(doc, "You will see: 'Account locked — too many failed attempts'")
    bullet(doc, "The lock lifts automatically after 15 minutes")
    bullet(doc, "A Super Admin can unlock your account immediately from User Management")
    h2(doc, "3.4  Logging Out")
    body(doc, "Always log out when finished, especially on shared computers.")
    step(doc, 1, "Find the logout option", "Look for your name or a logout button in the top-right corner of the portal.")
    step(doc, 2, "Click 'Logout'", "Your session ends and you return to the login page.")
    note_box(doc, "NOTE", "Your session expires automatically after a period of inactivity. You may be asked to log in again.")
    pb(doc)

    # ── SECTION 4 ─────────────────────────────────────────────────────────
    h1(doc, "4.  The Dashboard — Your Home Screen")
    body(doc, "The Dashboard is the first screen you see after logging in. It gives you a real-time overview of the entire platform's health and activity.")
    h2(doc, "4.1  Dashboard Elements Explained")
    table(doc, ["Element", "What It Shows", "Why It Matters"], [
        ["Total Vendors",           "Total businesses in the system",              "Tracks overall data collection progress"],
        ["Approved Vendors",        "Businesses that passed QA review",            "Shows data ready for the AirAd app"],
        ["Pending Vendors",         "Businesses awaiting QA review",               "Indicates backlog for the QA team"],
        ["Vendors Pending QA",      "Businesses flagged as 'Needs Review'",        "Urgent items requiring attention"],
        ["Vendors Approved Today",  "Approvals made today",                        "Daily productivity metric"],
        ["Imports Processing",      "CSV uploads being processed right now",       "Shows active background jobs"],
        ["QC Status Breakdown",     "Pie chart of vendor statuses",                "Visual overview of data quality"],
        ["14-Day Vendor Trend",     "Line chart of new vendors added daily",       "Shows data collection momentum"],
        ["7-Day Import Activity",   "Bar chart of CSV imports per day",            "Tracks bulk upload activity"],
        ["Recent Activity Feed",    "Last 10 actions by any user",                "Quick audit trail at a glance"],
    ])
    note_box(doc, "NOTE", "The dashboard is read-only. You cannot make changes from this screen. Use the navigation menu to go to the relevant section.")
    pb(doc)
