"""Shared helpers for AirAd User Guide generator."""
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _shd(cell, hex_color):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color)
    pr.append(s)


def h1(doc, txt):
    p = doc.add_heading(txt, level=1)
    for r in p.runs:
        r.font.color.rgb = RGBColor(21, 67, 96)


def h2(doc, txt):
    p = doc.add_heading(txt, level=2)
    for r in p.runs:
        r.font.color.rgb = RGBColor(40, 116, 166)


def body(doc, txt):
    p = doc.add_paragraph(txt)
    if p.runs:
        p.runs[0].font.size = Pt(10.5)


def bullet(doc, txt, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
        r.font.size = Pt(10.5)
    p.add_run(txt).font.size = Pt(10.5)


def note_box(doc, kind, txt):
    cfg = {
        "NOTE":      ("D6EAF8", (21, 67, 96)),
        "TIP":       ("D5F5E3", (11, 83, 69)),
        "WARNING":   ("FDEBD0", (120, 40, 31)),
        "IMPORTANT": ("F9EBEA", (100, 30, 22)),
    }
    bg, col = cfg.get(kind, cfg["NOTE"])
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.cell(0, 0)
    _shd(c, bg)
    p = c.paragraphs[0]
    r1 = p.add_run(f"{kind}: ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(*col)
    r2 = p.add_run(txt)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(*col)
    doc.add_paragraph()


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        _shd(c, "1F618D")
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        bg = "EBF5FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            _shd(c, bg)
            c.paragraphs[0].add_run(str(val)).font.size = Pt(9.5)
    doc.add_paragraph()


def step(doc, n, title, desc, subs=None):
    p = doc.add_paragraph()
    r1 = p.add_run(f"Step {n}: ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(21, 67, 96)
    r1.font.size = Pt(11)
    p.add_run(title).font.size = Pt(11)
    d = doc.add_paragraph(desc)
    d.paragraph_format.left_indent = Inches(0.3)
    if d.runs:
        d.runs[0].font.size = Pt(10)
    if subs:
        for s in subs:
            sp = doc.add_paragraph(style='List Bullet')
            sp.paragraph_format.left_indent = Inches(0.5)
            sp.add_run(s).font.size = Pt(10)


def pb(doc):
    doc.add_page_break()
