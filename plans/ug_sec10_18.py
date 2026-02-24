"""Sections 10-18: Tags, Governance, Audit, Users, Statuses, Quick Ref, FAQ, Glossary, Privacy"""
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ug_helpers import h1, h2, body, bullet, note_box, table, step, pb


def build(doc):

    # ── SECTION 10 ────────────────────────────────────────────────────────
    h1(doc, "10.  Tags — Organising & Categorising Vendors")
    body(doc, "Tags are labels that you attach to vendors to organise and categorise them. Tags make it easy to search, filter, and group vendors on both the portal and the AirAd app.")
    h2(doc, "10.1  Tag Types")
    table(doc, ["Tag Type", "Purpose", "Example"], [
        ["CATEGORY",   "Primary business category — required for approval",  "Restaurant, Pharmacy, Supermarket"],
        ["CUISINE",    "Type of food served (for food businesses)",           "Pakistani, Chinese, Fast Food"],
        ["FEATURE",    "Special features or amenities",                       "WiFi, Parking, Delivery"],
        ["PROMOTION",  "Active offers or deals",                              "50% Off, Happy Hour, New Opening"],
        ["SYSTEM",     "Auto-managed by the system — do not edit manually",  "verified, imported, flagged"],
    ])
    note_box(doc, "IMPORTANT", "SYSTEM tags are managed automatically by the platform. Do not manually add or remove SYSTEM tags — this may cause unexpected behaviour.")
    h2(doc, "10.2  Who Can Manage Tags?")
    body(doc, "Roles that can create and assign tags: Super Admin, City Manager, Data Entry, Data Quality Analyst.")
    h2(doc, "10.3  How to Create a New Tag")
    step(doc, 1, "Go to Tags", "Click 'Tags' in the left navigation menu.")
    step(doc, 2, "Click 'Add Tag'", "Find and click the Add Tag or '+' button.")
    step(doc, 3, "Fill in the details", "Enter the required information:", subs=[
        "Name: The display name of the tag (e.g. 'Restaurant')",
        "Slug: Lowercase, hyphen-separated identifier (e.g. 'restaurant')",
        "Tag Type: Select from CATEGORY, CUISINE, FEATURE, or PROMOTION",
        "Active: Tick to make the tag available for assignment",
    ])
    step(doc, 4, "Save", "Click Save. The tag is now available to assign to vendors.")
    h2(doc, "10.4  How to Assign a Tag to a Vendor")
    step(doc, 1, "Open the vendor", "Find and open the vendor detail page.")
    step(doc, 2, "Go to the Tags section", "Scroll to the Tags section or click the Tags tab.")
    step(doc, 3, "Search for a tag", "Type the tag name in the search box.")
    step(doc, 4, "Select the tag", "Click on the tag to assign it.")
    step(doc, 5, "Save", "Click Save. The tag is now linked to the vendor.")
    note_box(doc, "NOTE", "A vendor can have a maximum of 15 tags. If you try to add more, the system will show an error.")
    h2(doc, "10.5  How to Remove a Tag from a Vendor")
    step(doc, 1, "Open the vendor", "Find and open the vendor detail page.")
    step(doc, 2, "Go to the Tags section", "Find the tag you want to remove.")
    step(doc, 3, "Click the remove icon", "Click the X or Remove button next to the tag.")
    step(doc, 4, "Save", "Click Save.")
    note_box(doc, "WARNING", "Removing the only CATEGORY tag from an APPROVED vendor will change its status back to NEEDS_REVIEW, as it no longer meets the approval requirement.")
    pb(doc)

    # ── SECTION 11 ────────────────────────────────────────────────────────
    h1(doc, "11.  Governance — Fraud, Blacklists & Enforcement")
    body(doc, "The Governance section handles platform integrity — detecting fraud, managing blacklists, suspending vendors, and ensuring compliance with terms of service and data privacy regulations.")
    body(doc, "Access to Governance is restricted to: Super Admin, Operations Manager, Content Moderator.")
    h2(doc, "11.1  Fraud Scores")
    body(doc, "Every vendor has a Fraud Score — a number between 0 and 100 that indicates the likelihood of the vendor data being fraudulent or inaccurate.")
    table(doc, ["Score Range", "Risk Level", "Recommended Action"], [
        ["0 – 30",   "Low Risk",    "No action needed — data appears legitimate"],
        ["31 – 60",  "Medium Risk", "Review the vendor data carefully before approving"],
        ["61 – 80",  "High Risk",   "Flag for QA review; investigate before approving"],
        ["81 – 100", "Critical",    "Consider blacklisting; do not approve without thorough investigation"],
    ])
    body(doc, "Fraud scores are calculated automatically based on factors such as duplicate GPS coordinates, suspicious phone numbers, and data inconsistencies.")
    h2(doc, "11.2  The Blacklist")
    body(doc, "The Blacklist is a list of vendors or entities that are permanently banned from the platform. Blacklisted vendors cannot be approved and are hidden from the app.")
    h2(doc, "11.3  How to Blacklist a Vendor")
    step(doc, 1, "Go to Governance", "Click 'Governance' in the left navigation menu.")
    step(doc, 2, "Find the vendor", "Search for the vendor you want to blacklist.")
    step(doc, 3, "Click 'Add to Blacklist'", "Click the Blacklist button.")
    step(doc, 4, "Enter the reason", "Provide a clear reason for blacklisting (e.g. 'Fraudulent data', 'Duplicate of existing vendor').")
    step(doc, 5, "Confirm", "Click Confirm. The vendor is blacklisted immediately.")
    note_box(doc, "WARNING", "Blacklisting is a serious action. It immediately hides the vendor from the app and prevents future approval. Only blacklist vendors when there is clear evidence of fraud or policy violation.")
    h2(doc, "11.4  Vendor Suspensions")
    body(doc, "A Suspension is a temporary enforcement action — less severe than blacklisting. A suspended vendor is hidden from the app for a defined period.")
    table(doc, ["Field", "Description"], [
        ["Reason",       "Why the vendor is being suspended"],
        ["Start Date",   "When the suspension begins"],
        ["End Date",     "When the suspension automatically lifts (optional)"],
        ["Suspended By", "Which admin user applied the suspension"],
    ])
    h2(doc, "11.5  How to Suspend a Vendor")
    step(doc, 1, "Go to Governance", "Click 'Governance' in the navigation.")
    step(doc, 2, "Find the vendor", "Search for the vendor.")
    step(doc, 3, "Click 'Suspend'", "Click the Suspend button.")
    step(doc, 4, "Fill in suspension details", "Enter the reason, start date, and optional end date.")
    step(doc, 5, "Confirm", "Click Confirm. The vendor is suspended.")
    h2(doc, "11.6  Terms of Service Acceptance")
    body(doc, "The system tracks whether vendors (or their representatives) have accepted the AirAd Terms of Service. This record is stored permanently and cannot be deleted.")
    h2(doc, "11.7  GDPR Consent Records")
    body(doc, "For vendors in regions covered by GDPR (General Data Protection Regulation), the system stores consent records — confirming that data was collected with proper consent. These records include:")
    bullet(doc, "The date and time consent was given")
    bullet(doc, "The version of the privacy policy accepted")
    bullet(doc, "The IP address from which consent was given")
    note_box(doc, "IMPORTANT", "GDPR consent records are immutable — they can never be edited or deleted. This is a legal requirement.")
    pb(doc)

    # ── SECTION 12 ────────────────────────────────────────────────────────
    h1(doc, "12.  Audit Log — Full Activity History")
    body(doc, "The Audit Log is a complete, tamper-proof record of every significant action taken in the portal. Every time someone creates, edits, or deletes data, an audit log entry is automatically created.")
    h2(doc, "12.1  Who Can View the Audit Log?")
    body(doc, "Roles that can view the audit log: Super Admin, Analyst.")
    h2(doc, "12.2  What Does an Audit Log Entry Show?")
    table(doc, ["Field", "Description", "Example"], [
        ["Timestamp",    "Exact date and time of the action",                    "2024-03-15 14:32:07"],
        ["User",         "Who performed the action",                             "ahmed@airaad.com"],
        ["Action",       "What type of action was taken",                        "UPDATE, CREATE, DELETE"],
        ["Resource Type","What type of data was changed",                        "Vendor, ImportBatch, AdminUser"],
        ["Resource ID",  "The unique ID of the record that was changed",         "a3f2-..."],
        ["Before",       "What the data looked like before the change",          "{'status': 'PENDING'}"],
        ["After",        "What the data looks like after the change",            "{'status': 'APPROVED'}"],
        ["IP Address",   "The IP address of the user who made the change",       "192.168.1.10"],
    ])
    h2(doc, "12.3  How to View the Audit Log")
    step(doc, 1, "Go to Audit Log", "Click 'Audit Log' in the left navigation menu.")
    step(doc, 2, "Browse or search", "You can filter by date range, user, action type, or resource type.")
    step(doc, 3, "Click an entry", "Click on any entry to see the full before/after snapshot of the change.")
    note_box(doc, "IMPORTANT", "Audit log entries are permanent and cannot be edited or deleted — not even by a Super Admin. This ensures a trustworthy record of all activity.")
    h2(doc, "12.4  Common Uses of the Audit Log")
    bullet(doc, "Find out who approved or rejected a specific vendor")
    bullet(doc, "See what changes were made to a vendor record and when")
    bullet(doc, "Investigate suspicious activity or data errors")
    bullet(doc, "Verify that a CSV import was processed correctly")
    bullet(doc, "Confirm that a user account was created or modified")
    pb(doc)

    # ── SECTION 13 ────────────────────────────────────────────────────────
    h1(doc, "13.  User Management (Super Admin Only)")
    body(doc, "User Management allows Super Admins to create, edit, and manage all portal user accounts. Only Super Admins have access to this section.")
    h2(doc, "13.1  How to Create a New User")
    step(doc, 1, "Go to User Management", "Click 'Users' in the left navigation menu.")
    step(doc, 2, "Click 'Add User'", "Find and click the Add User or '+' button.")
    step(doc, 3, "Fill in user details", "Enter the required information:", subs=[
        "Full Name: The user's full name",
        "Email Address: Their work email (used as login username)",
        "Role: Select the appropriate role from the dropdown",
        "Temporary Password: Set an initial password — the user must change it on first login",
    ])
    step(doc, 4, "Save", "Click Save. The user account is created and they can now log in.")
    note_box(doc, "TIP", "Inform the new user of their email and temporary password via a secure channel (e.g. in person or via encrypted message). Do not send passwords by email.")
    h2(doc, "13.2  How to Edit a User")
    step(doc, 1, "Go to User Management", "Click 'Users' in the navigation.")
    step(doc, 2, "Find the user", "Search for the user by name or email.")
    step(doc, 3, "Click on the user", "Click the user's name to open their profile.")
    step(doc, 4, "Click 'Edit'", "Click the Edit button.")
    step(doc, 5, "Make changes", "Update the name, role, or other details as needed.")
    step(doc, 6, "Save", "Click Save.")
    h2(doc, "13.3  How to Unlock a Locked Account")
    step(doc, 1, "Go to User Management", "Click 'Users' in the navigation.")
    step(doc, 2, "Find the locked user", "Locked accounts are marked with a 'Locked' badge.")
    step(doc, 3, "Click 'Unlock'", "Click the Unlock button next to the user.")
    step(doc, 4, "Confirm", "The account is unlocked immediately and the user can log in again.")
    h2(doc, "13.4  GDPR — Exporting a User's Data")
    body(doc, "Under GDPR, users have the right to request a copy of all data held about them. Super Admins can export this data:")
    step(doc, 1, "Go to User Management", "Find the user.")
    step(doc, 2, "Click 'Export Data'", "Click the GDPR Export or Export Data button.")
    step(doc, 3, "Download", "A JSON file containing all data associated with that user is downloaded.")
    h2(doc, "13.5  GDPR — Deleting a User's Data")
    body(doc, "Users can request deletion of their personal data (Right to be Forgotten). Super Admins can action this:")
    step(doc, 1, "Go to User Management", "Find the user.")
    step(doc, 2, "Click 'Delete Personal Data'", "Click the GDPR Delete or Anonymise button.")
    step(doc, 3, "Confirm", "The user's personal data is anonymised. The account record remains for audit purposes but personal details are removed.")
    note_box(doc, "IMPORTANT", "GDPR deletion anonymises the user's personal data but does NOT delete their audit log entries. Audit logs are permanent by law.")
    pb(doc)

    # ── SECTION 14 ────────────────────────────────────────────────────────
    h1(doc, "14.  Understanding Statuses at a Glance")
    h2(doc, "14.1  Vendor QC Statuses")
    table(doc, ["Status", "Colour", "Meaning", "Visible on App?"], [
        ["PENDING",      "Grey",   "Newly added, not yet reviewed",                    "No"],
        ["APPROVED",     "Green",  "Passed QA review — live on the app",               "Yes"],
        ["REJECTED",     "Red",    "Failed QA review",                                 "No"],
        ["NEEDS_REVIEW", "Orange", "Flagged for re-inspection (GPS drift, duplicate)", "No"],
        ["FLAGGED",      "Yellow", "Manually flagged for attention",                   "No"],
    ])
    h2(doc, "14.2  Import Batch Statuses")
    body(doc, "These statuses apply to both CSV uploads and Google Places seed jobs:")
    table(doc, ["Status", "Shown As", "Meaning"], [
        ["QUEUED",      "Pending",     "The import has been received and is waiting to start processing"],
        ["PROCESSING",  "In progress", "The system is actively processing the import right now"],
        ["DONE",        "Done",        "All rows or businesses have been processed successfully"],
        ["FAILED",      "Failed",      "The import failed — the system retries automatically up to 3 times"],
    ])
    h2(doc, "14.3  Field Visit Outcomes")
    table(doc, ["Outcome", "Meaning"], [
        ["Verified",     "Business confirmed to exist and data is accurate"],
        ["Not Found",    "No business found at the given location"],
        ["Closed Down",  "Business has permanently closed"],
        ["Needs Update", "Business exists but data needs correction"],
    ])
    pb(doc)

    # ── SECTION 15 ────────────────────────────────────────────────────────
    h1(doc, "15.  Common Tasks — Quick Reference")
    table(doc, ["Task", "Where to Go", "Who Can Do It"], [
        ["Add a single vendor",           "Vendors → Add Vendor",                  "Super Admin, City Manager, Data Entry"],
        ["Import vendors from CSV",       "Imports → CSV upload tab",              "Super Admin, City Manager, Data Entry"],
        ["Seed vendors from Google Places","Imports → Google Places tab",           "Super Admin, City Manager, Operations Manager"],
        ["Approve a vendor",              "QA → Select vendor → Approve",          "Super Admin, City Manager, QA Reviewer"],
        ["Reject a vendor",               "QA → Select vendor → Reject",           "Super Admin, City Manager, QA Reviewer"],
        ["Flag a vendor for review",      "Vendor Detail → Flag for Review",       "Super Admin, City Manager, QA Reviewer"],
        ["Add a city",                    "Geo → Add City",                        "Super Admin, City Manager, Data Entry"],
        ["Add an area",                   "Geo → Add Area",                        "Super Admin, City Manager, Data Entry"],
        ["Record a field visit",          "Field Ops → New Visit",                 "Super Admin, City Manager, Field Agent"],
        ["Upload photos",                 "Field Ops → Visit → Add Photos",        "Super Admin, City Manager, Field Agent"],
        ["Create a tag",                  "Tags → Add Tag",                        "Super Admin, City Manager, Data Entry"],
        ["Assign a tag to a vendor",      "Vendor Detail → Tags → Add Tag",        "Super Admin, City Manager, Data Entry"],
        ["View audit log",                "Audit Log",                             "Super Admin, Analyst"],
        ["Create a new user",             "Users → Add User",                      "Super Admin only"],
        ["Unlock a locked account",       "Users → Find User → Unlock",           "Super Admin only"],
        ["Blacklist a vendor",            "Governance → Blacklist",                "Super Admin, Operations Manager"],
        ["Suspend a vendor",              "Governance → Suspend",                  "Super Admin, Operations Manager, Content Moderator"],
        ["Export user data (GDPR)",       "Users → Find User → Export Data",       "Super Admin only"],
        ["Check import errors",           "Imports → Click import → View errors",  "Super Admin, City Manager, Data Entry"],
        ["Check Google Places progress",  "Imports → Google Places tab → Batches table", "Super Admin, City Manager, Operations Manager"],
    ])
    pb(doc)

    # ── SECTION 16 ────────────────────────────────────────────────────────
    h1(doc, "16.  Troubleshooting & FAQs")
    h2(doc, "16.1  Login & Account Issues")
    table(doc, ["Problem", "Likely Cause", "Solution"], [
        ["Cannot log in — wrong password",        "Incorrect password entered",                    "Check Caps Lock is off; try again carefully"],
        ["Account is locked",                     "5+ failed login attempts",                      "Wait 15 minutes or ask Super Admin to unlock"],
        ["Forgot password",                       "Password not remembered",                       "Ask your Super Admin to reset your password"],
        ["Page not loading",                      "Browser or network issue",                      "Try refreshing; try a different browser; check internet connection"],
        ["Session expired",                       "Inactive for too long",                         "Log in again — your work is saved"],
    ])
    h2(doc, "16.2  Vendor Issues")
    table(doc, ["Problem", "Likely Cause", "Solution"], [
        ["Cannot find a vendor",                  "Wrong search term or filter",                   "Clear all filters and search by exact business name"],
        ["Vendor not showing on app",             "Vendor is not APPROVED",                        "Check vendor status — it must be APPROVED to appear on app"],
        ["Cannot approve vendor",                 "No CATEGORY tag assigned",                      "Assign at least one Category tag first, then approve"],
        ["Duplicate vendor warning",              "Similar vendor already exists",                 "Check if the business is already in the system; reject the duplicate"],
        ["GPS coordinates look wrong",            "Data entry error",                              "Edit the vendor and enter correct latitude/longitude"],
    ])
    h2(doc, "16.3  CSV Import Issues")
    table(doc, ["Problem", "Likely Cause", "Solution"], [
        ["Import failed immediately",             "Wrong file format or missing required columns", "Check file is .csv and has all required column headers"],
        ["Many rows failed with city_slug error", "City slugs in CSV do not match system",         "Go to Geo section, copy exact slugs, update CSV and re-upload"],
        ["Import stuck on PROCESSING",            "Large file or system load",                     "Wait up to 10 minutes; if still stuck, contact Super Admin"],
        ["Duplicate vendor errors",               "Businesses already exist in system",            "Remove duplicate rows from CSV or check if they are already imported"],
    ])
    h2(doc, "16.4  Frequently Asked Questions")
    for q, a in [
        ("Can I undo a deletion?",
         "Vendors are soft-deleted (hidden, not permanently removed). A Super Admin can restore a soft-deleted vendor by contacting the technical team."),
        ("Why can I not see the Governance section?",
         "Governance is only accessible to Super Admin, Operations Manager, and Content Moderator. If you need access, ask your Super Admin."),
        ("Can two users have the same email address?",
         "No. Each user account must have a unique email address."),
        ("How do I know which city slug to use in my CSV?",
         "Go to the Geo section in the portal and look at the slug column for each city. Copy it exactly."),
        ("What happens if I upload the same CSV twice?",
         "The system will attempt to process it again. Rows that already exist may be flagged as duplicates. Always check before re-uploading."),
        ("What is Google Places Seeding?",
         "It is a way to automatically pull business listings from Google into AirAd. You select an area and keywords, and the system finds and imports matching businesses. See Section 7.8 for full details."),
        ("Why do I not see the Google Places tab?",
         "Only Super Admin, City Manager, and Operations Manager roles can use Google Places Seeding. If you are a Data Entry user, you only see CSV Upload. Contact your Super Admin if you need access."),
        ("Can a Field Agent approve vendors?",
         "No. Field Agents can only record visits and upload photos. Approval is done by QA Reviewers, City Managers, or Super Admins."),
        ("How long are photo links valid?",
         "Photo links are temporary and expire after a short time for security. To view a photo again, open the vendor or visit record in the portal."),
        ("Can I export vendor data to Excel?",
         "This depends on your portal configuration. Ask your Super Admin if a data export feature is available for your role."),
    ]:
        p = doc.add_paragraph()
        r = p.add_run(f"Q: {q}")
        r.bold = True; r.font.size = Pt(10.5)
        doc.add_paragraph(f"A: {a}").runs[0].font.size = Pt(10.5)
        doc.add_paragraph()
    pb(doc)

    # ── SECTION 17 ────────────────────────────────────────────────────────
    h1(doc, "17.  Glossary of Terms")
    body(doc, "This glossary explains technical and portal-specific terms used throughout this guide.")
    table(doc, ["Term", "Definition"], [
        ["Vendor",          "Any business or point of interest registered on the AirAd platform"],
        ["Slug",            "A short, URL-friendly identifier using lowercase letters and hyphens (e.g. 'karachi', 'al-fatah-superstore')"],
        ["QC Status",       "Quality Control Status — the review state of a vendor (PENDING, APPROVED, REJECTED, NEEDS_REVIEW, FLAGGED)"],
        ["CSV",             "Comma-Separated Values — a simple spreadsheet file format used for bulk data imports"],
        ["GPS",             "Global Positioning System — technology used to determine exact geographic coordinates (latitude and longitude)"],
        ["Latitude",        "The north-south position of a location on Earth (e.g. 24.8607 for Karachi)"],
        ["Longitude",       "The east-west position of a location on Earth (e.g. 67.0011 for Karachi)"],
        ["Geo Hierarchy",   "The geographic structure: Country → City → Area → Landmark"],
        ["Soft Delete",     "Hiding a record from the app without permanently removing it from the database"],
        ["Audit Log",       "A permanent, tamper-proof record of every action taken in the portal"],
        ["Field Visit",     "An on-site visit by a Field Agent to verify a vendor's physical location"],
        ["GPS Drift",       "A significant difference between stored GPS coordinates and field-confirmed coordinates"],
        ["Fraud Score",     "A 0-100 score indicating the likelihood of vendor data being fraudulent"],
        ["Blacklist",       "A permanent ban preventing a vendor from appearing on the platform"],
        ["Suspension",      "A temporary ban preventing a vendor from appearing on the platform"],
        ["GDPR",            "General Data Protection Regulation — EU law governing personal data collection and storage"],
        ["AES-256-GCM",     "A strong encryption standard used to protect sensitive data like phone numbers"],
        ["Celery",          "Background task system used to process CSV imports without slowing the portal"],
        ["S3",              "Amazon S3 — secure cloud storage where uploaded files and photos are stored"],
        ["Role",            "A set of permissions assigned to a portal user that controls what they can see and do"],
        ["Tag",             "A label attached to a vendor to categorise or describe it"],
        ["Import Batch",    "A single import job — either a CSV upload or a Google Places seed request"],
        ["Google Places Seeding", "An automated way to import business listings from Google's database by selecting a geographic area and search keywords"],
        ["Search Radius",  "The distance in metres around the centre of an area used when searching for businesses on Google Places"],
        ["Category Tag",   "A label describing the type of business (e.g. Restaurant, Pharmacy) — used to organise and filter vendors"],
        ["AR",              "Augmented Reality — the technology used in the AirAd mobile app to overlay business info on camera view"],
    ])
    pb(doc)

    # ── SECTION 18 ────────────────────────────────────────────────────────
    h1(doc, "18.  Data Privacy & Security")
    body(doc, "The AirAd Data Collection Portal is built with data privacy and security as a top priority. This section explains how your data and the data you collect is protected.")
    h2(doc, "18.1  How Your Account is Protected")
    bullet(doc, "Your password is never stored in plain text — it is securely hashed", bold_prefix="Password Security")
    bullet(doc, "Your account locks after 5 failed login attempts to prevent unauthorised access", bold_prefix="Account Lockout")
    bullet(doc, "Your login session expires automatically after a period of inactivity", bold_prefix="Session Expiry")
    bullet(doc, "All data is transmitted over HTTPS — encrypted in transit", bold_prefix="Encrypted Connection")
    h2(doc, "18.2  How Vendor Data is Protected")
    bullet(doc, "Phone numbers are encrypted using AES-256-GCM encryption — the same standard used by banks", bold_prefix="Phone Number Encryption")
    bullet(doc, "Photos are stored in secure cloud storage (Amazon S3) with access-controlled temporary links", bold_prefix="Photo Storage")
    bullet(doc, "Deleted vendors are soft-deleted — hidden but preserved for audit and legal purposes", bold_prefix="Soft Deletion")
    bullet(doc, "Every change to vendor data is recorded in the immutable Audit Log", bold_prefix="Audit Trail")
    h2(doc, "18.3  GDPR Compliance")
    body(doc, "The portal is designed to comply with GDPR requirements:")
    bullet(doc, "Consent records are stored permanently and cannot be altered")
    bullet(doc, "Users can request export of all data held about them (Right of Access)")
    bullet(doc, "Users can request anonymisation of their personal data (Right to be Forgotten)")
    bullet(doc, "All data processing activities are logged in the Audit Log")
    h2(doc, "18.4  Your Responsibilities")
    note_box(doc, "IMPORTANT", "As a portal user, you are responsible for the data you enter. Always ensure the data you collect is accurate, lawful, and obtained with proper consent.")
    bullet(doc, "Only enter data you have permission to collect")
    bullet(doc, "Do not enter personal data (e.g. individual names, personal phone numbers) unless required")
    bullet(doc, "Do not share your login credentials with anyone")
    bullet(doc, "Report any suspected data breach or suspicious activity to your Super Admin immediately")
    bullet(doc, "Log out of the portal when you are finished, especially on shared devices")
    h2(doc, "18.5  Reporting a Security Issue")
    body(doc, "If you notice anything suspicious — such as data being changed without your knowledge, or someone else using your account — report it immediately to your Super Admin. Do not attempt to investigate or fix security issues yourself.")
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("— End of AirAd Data Collection Portal User Guide —")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(21, 67, 96)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Version 1.0  |  For internal use only  |  AirAd Platform")
    r2.font.size = Pt(10)
